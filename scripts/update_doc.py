from docx import Document

doc = Document('ReconPilot_Complete_Implementation_Guide 1 (1).docx')

# 1. Update Headings
for p in doc.paragraphs:
    if "5. Database and Audit Model" in p.text:
        p.text = p.text.replace("5. Database and Audit Model", "5. Audit Model (Database is Planned Future Work)")
    elif "10.2 Approval gates" in p.text:
        p.text = p.text.replace("10.2 Approval gates", "10.2 Approval gates (Design — Not Yet Implemented)")
    elif "14.2 Dockerfile" in p.text:
        p.text = p.text.replace("14.2 Dockerfile", "14.2 Dockerfile (Planned for future packaging)")
    elif "9.1 Connection steps" in p.text:
        # Add mock mode note after this section
        pass

# Add a note under Section 9.1
for i, p in enumerate(doc.paragraphs):
    if "9.1 Connection steps" in p.text:
        # Insert a new paragraph after the list (around index i+8)
        pass # Better to append the note about mock mode at the end of section 9

# Let's add the Forward Cash Forecaster section at the end, before Section 16
# Actually, it's easier to append to the document.
doc.add_heading('17. Forward Cash Forecaster (New Feature)', level=1)
doc.add_paragraph('ReconPilot includes a deterministic forward cash forecaster that projects expected cash inflows over a configurable horizon based on already-computed reconciliation results and historical settlement timing. This satisfies the "Forward cash forecaster" example direction from the Track 04 problem statement.')
doc.add_paragraph('The forecaster uses pure arithmetic, avoiding LLM hallucinations for financial projections. It categorizes records into settled, expected (using the SLA window for missing settlements), at-risk, and excluded, giving a clear, actionable projection of future cash flows.')

# 2. Update Tables
for i, table in enumerate(doc.tables):
    if not table.rows:
        continue
    first_cell_text = table.rows[0].cells[0].text
    
    # Table 6: Repo Structure
    if "reconpilot/" in first_cell_text and ".env.example" in first_cell_text:
        table.rows[0].cells[0].text = """reconpilot/
|-- .env.example
|-- .gitignore
|-- README.md
|-- requirements.txt
|-- backend/
|   |-- main.py
|   |-- config.py
|   |-- razorpay_client.py
|   |-- validation.py
|   |-- normalization.py
|   |-- matcher.py
|   |-- classifier.py
|   |-- pipeline.py
|   |-- audit.py
|   |-- forecaster.py
|   |-- agent_tools.py
|   `-- agent.py
|-- frontend/
|   `-- streamlit_app.py
|-- scripts/
|   |-- generate_synthetic_data.py
|   |-- evaluate.py
|   `-- scenario_manifest.csv
`-- tests/
    |-- test_validation.py
    |-- test_matching.py
    |-- test_classifier.py
    `-- test_api_failure.py"""

    # Table 12: Config
    elif "# backend/config.py" in first_cell_text:
        table.rows[0].cells[0].text = """# backend/config.py
from pydantic_settings import BaseSettings, SettingsConfigDict

class Settings(BaseSettings):
    razorpay_key_id: str = ""
    razorpay_key_secret: str = ""
    razorpay_base_url: str = "https://api.razorpay.com/v1"
    llm_api_key: str = ""
    llm_model: str = "claude-sonnet-4-6"
    api_timeout_seconds: int = 20
    max_api_attempts: int = 3
    amount_tolerance_paise: int = 0
    probable_match_threshold_days: int = 3
    settlement_sla_days: int = 5
    model_config = SettingsConfigDict(env_file=".env", extra="ignore")

settings = Settings()"""

    # Table 15: Generator logic
    elif "# scripts/generate_synthetic_data.py" in first_cell_text:
        table.rows[0].cells[0].text = """# scripts/generate_synthetic_data.py generates 9 distinct scenarios (exact match, refund-adjusted, fee/tax mismatches, missing settlements, duplicates, amount mismatches, delayed settlements, validation errors, and ambiguous composites) across 4 CSV files (orders, transactions, bank settlements, refunds).
# It also generates a separate scenario_manifest.csv for evaluation."""

    # Table 24: Agent tools
    elif "get_batch_summary(batch_id)" in first_cell_text:
        table.rows[0].cells[0].text = """get_batch_summary(batch_id)
get_record_details(record_id)
explain_matching_rule(record_id)
list_exceptions(batch_id, exception_type=None)
get_unresolved_value(batch_id)
get_cash_forecast(batch_id, horizon_days)"""

    # Table 31: API Endpoints
    elif "POST /batches" in first_cell_text:
        table.rows[0].cells[0].text = """POST /batches/run          Run validation and reconciliation
POST /agent/query          Ask a bounded question
GET  /health               Health check
GET  /razorpay/status      Razorpay mock mode status
(Other endpoints are planned for future extension)"""

    # Table 32: FastAPI main
    elif "# backend/main.py" in first_cell_text:
        table.rows[0].cells[0].text = """# backend/main.py
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from . import agent as agent_module, razorpay_client
from .pipeline import run_batch

app = FastAPI(title="ReconPilot API", version="1.0.0")

@app.post("/batches/run")
def run(batch_id: str = None):
    return run_batch(batch_id=batch_id)

@app.post("/agent/query")
def agent_query(payload: AgentQuery):
    return agent_module.ask(payload.question, payload.batch_id)"""

    # Table 33: Streamlit pages
    elif "Page" in first_cell_text and len(table.rows) > 1 and "Upload and Run" in table.rows[1].cells[0].text:
        # Clear existing rows except header
        for _ in range(len(table.rows) - 1):
            table._tbl.remove(table.rows[1]._tr)
        
        # Add correct tabs
        row_cells = table.add_row().cells
        row_cells[0].text, row_cells[1].text = "Overview", "Run a batch, metrics cards, status distribution, delayed settlement warning"
        row_cells = table.add_row().cells
        row_cells[0].text, row_cells[1].text = "Exceptions", "Filter exceptions table by exception type"
        row_cells = table.add_row().cells
        row_cells[0].text, row_cells[1].text = "Ask ReconPilot", "Prompt box plus evidence-backed answer and cited record IDs"
        row_cells = table.add_row().cells
        row_cells[0].text, row_cells[1].text = "Cash Forecast", "Forward cash projections based on reconciliation results"

    # Table 35: Test example
    elif "def test_amount_mismatch_is_not_auto_matched" in first_cell_text:
        table.rows[0].cells[0].text = """def test_amount_mismatch_is_not_auto_matched():
    bank = base_bank()
    bank["credited_amount_paise"] += 100
    decision = reconcile(base_order(), [base_payment()], [bank], [])
    assert decision.status == "AMOUNT_MISMATCH"
    assert decision.requires_review is True
    assert decision.evidence["actual_credit_paise"] != decision.evidence["expected_net_paise"]"""


doc.save('ReconPilot_Complete_Implementation_Guide 1 (1).docx')
print("Document updated successfully.")
